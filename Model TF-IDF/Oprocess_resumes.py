# -*- coding: utf-8 -*-
"""
Created on Sun Jun  3 18:41:34 2018

@author: prasasaw
TF-IDF and Cosine similarity are generally used in combination with each other for finding document similarity. TF-IDF helps to find a subset of words that distinguish a document rather than using all words in that document.  The TF-IDF vectors for 2 documents are then compared using Cosine similarity index to see how similar the documents are.

Universities : NNP parsing of school, college and universities

Secondary:
Experience (high)
certi
Companies
Education

See code examples from Sci kit
http://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html#sklearn.feature_extraction.text.TfidfVectorizer.build_preprocessor

try spacy
"""
import os
import docx
from docx import Document
import pandas as pd
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
import re
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
import glob
import datefinder
from datetime import date
from datetime import datetime
import yaml
import lib
import numpy as np
import gensim
from gensim import corpora
from gensim.models import word2vec
import func
from nltk.tokenize import TreebankWordTokenizer
from gensim import corpora, models
import pickle
import PyPDF2
from nltk.tokenize import word_tokenize


dir =  r"C:\Resumes\WORD resumes"
files = glob.glob(os.path.join(dir,"*"))  

filepath =[]
Employee = []
FullText = []
CleanseText = []
Texts = []
birthDate = []
ages = []
email = []
phone = []
programming = []
packages = []
domains = []
roles = []

## Extract Text from Documents
for filename in files:
    filepath.append(filename)
    fileext = os.path.splitext(filename)[1]
    newname = os.path.splitext(filename)[0]
    newname = os.path.split(newname)[1]
    # cleanse Profile Name
    newname = func.cleanse_name(newname)
    OneText = ''
    
    regexdoc = re.compile('.*doc.*')
    regexpdf = re.compile('.*pdf.*')
    
    matchdoc = regexdoc.search(fileext)
    matchpdf = regexpdf.search(fileext)
    
    if matchdoc:    
        wordDoc = Document(filename)
       
        # Extract all text from paragraphs and tables on OneText
                     
            
        for para in wordDoc.paragraphs:
            OneText = OneText + '\n' + para.text
            
        for table in wordDoc.tables:
            rowText = ''
            for row in table.rows:
                rowCell = ''
                for cell in row.cells:
                    rowCell = rowCell.rstrip()
                    rowCell = rowCell + cell.text + ':' 
                rowText = rowText + rowCell + '\n'
            
            OneText = OneText + rowText  + '\n'
    
    if matchpdf:
        pdffile=open(filename, 'rb')
        pdfreader=PyPDF2.PdfFileReader(pdffile)
                
        content = ''
        for i in range(0, pdfreader.getNumPages()):
            # Extract text from page and add to content
            content += pdfreader.getPage(i).extractText() + "\n"
        OneText = content   
    
    #check if onetext is blank, then break
    empty = []
    if OneText == '':
        empty.append(newname)
        
    
    
    tokens = func.process_message(OneText, stem = False)
    Cleanse_Text = str(tokens)
    Texts.append(tokens)
    CleanseText.append(Cleanse_Text)
    
    Employee.append(newname)
    FullText.append(OneText)
    #--end
    
    # extract birth date
    dob = func.find_BirthDate(OneText)
    birthDate.append(dob)
    
    # calculate age
    age = ''
    if birthDate[0] != '':    
        age=func.calculate_age(dob)
    ages.append(age)       
    
    # extract email
    emailid = ''
    emailid = func.extract_email(OneText)
    email.append(emailid)
    
    # extract phone
    ph = ''
    ph = func.extract_phone(OneText)
    phone.append(ph)
    
    
    # extract skills
    for categories, x in func.get_conf('Categories').items():
        
        if categories == 'programming':
            b = func.extract_skills(OneText,categories,x)
            programming.append(b)
    
        if categories == 'packages':
            b = func.extract_skills(OneText,categories,x)
            packages.append(b)
       
        if categories == 'domain':
            b = func.extract_skills(OneText,categories,x)
            domains.append(b)

        if categories == 'roles':
            b = func.extract_skills(OneText,categories,x)
            roles.append(b)


df = pd.DataFrame({'Employee' : Employee, 'FullText' : FullText, 'DOB' : birthDate,
                   'Email' : email, 'Age' : ages, 'Programming' : programming,
                   'Packages' : packages, 'Domains' : domains, 'Roles' : roles,
                   'Phone' : phone, 'CleanseText' : CleanseText, 'Filepath': filepath}, 
                  columns=['Employee', 'FullText', 'CleanseText', 'Age', 'Email', 'Phone', 'Programming', 'Packages', 'Domains', 'Roles', 'Filepath'])

df['CleanseText'] = df['CleanseText'].str.replace('[,\[\]"\']+', ' ')

# TF IDF vectorizer
# create a dictionary
dictionary = corpora.Dictionary(Texts)
# convert tokenized documents into a document-term matrix
corpus = [dictionary.doc2bow(Text) for Text in Texts]
tf_idf = gensim.models.TfidfModel(corpus)

# Now we will create a similarity measure object in tf-idf space.
# "Tell me how similar is this query document to each document in the index?"
sims = gensim.similarities.Similarity('index/temp.index',tf_idf[corpus],
                                      num_features=len(dictionary))

#Now create a query document and convert it to tf-idf
query_doc = [w.lower() for w in word_tokenize("TOGAF and data science master data management NLP Python")]
print(query_doc)
query_doc_bow = dictionary.doc2bow(query_doc)
print(query_doc_bow)
#gives the  TF-IDF score for each term in the query. refer https://janav.wordpress.com/2013/10/27/tf-idf-and-cosine-similarity/
query_doc_tf_idf = tf_idf[query_doc_bow]
print(query_doc_tf_idf)

result=sims[query_doc_tf_idf]
