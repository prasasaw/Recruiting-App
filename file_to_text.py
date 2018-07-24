# -*- coding: utf-8 -*-
"""
Created on Fri Jul 13 10:43:46 2018
@author: prasasaw

input: single file input
output: text string

TODO:
filepath =[]
Employee = []

"""
import os
import re
import docx
from docx import Document
import func

def text_converter(filename):

    """dir =  r"C:\Resume Miner1"
    files = glob.glob(os.path.join(dir,"*"))  
    
    for filename in files:"""

    fileext = os.path.splitext(filename)[1]
    OneText = ''
    
    regexdoc = re.compile('.*doc.*')
    regexpdf = re.compile('.*pdf.*')
    
    matchdoc = regexdoc.search(fileext)
    matchpdf = regexpdf.search(fileext)
    
    if matchdoc:    
        wordDoc = Document(filename)
       
        # Extract all text from paragraphs and tables on OneText
        for para in wordDoc.paragraphs:
            OneText = OneText + '\n' + para.text
            
        for table in wordDoc.tables:
            rowText = ''
            for row in table.rows:
                rowCell = ''
                for cell in row.cells:
                    rowCell = rowCell.rstrip()
                    rowCell = rowCell + cell.text + ':' 
                rowText = rowText + rowCell + '\n'
            
            OneText = OneText + rowText  + '\n'
    
    if matchpdf:
        OneText = func.extract_text_from_pdf(filename)  
    
    return(OneText)
        